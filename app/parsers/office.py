import asyncio
from pathlib import Path

from docx import Document as DocxDocument
from langchain_community.document_loaders import UnstructuredPowerPointLoader
from langchain_core.callbacks import adispatch_custom_event
from langchain_core.documents import Document
from openpyxl import load_workbook

from app.parsers.base import FileParser


class ExcelParser(FileParser):
    async def parse(self, file_path: Path) -> list[Document]:
        """
        将 Excel 每一行转换为带上下文的 Document 对象。
        """
        # 加载工作簿是阻塞操作
        wb = await asyncio.to_thread(load_workbook, str(file_path), data_only=True)
        documents = []

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            # 获取表头
            headers = [str(cell.value) if cell.value is not None else "" for cell in sheet[1]]

            rows = list(sheet.iter_rows(min_row=2))
            total_rows = len(rows)

            for row_idx, row in enumerate(rows, start=2):
                if row_idx % 100 == 0:
                    await adispatch_custom_event(
                        "status_update", {"text": f"正在解析 Excel ({sheet_name}): {row_idx}/{total_rows} 行..."}
                    )

                # 行处理也可能耗时，在大规模数据下可考虑分批 thread 执行，这里保持简单
                row_values = [str(cell.value) if cell.value is not None else "" for cell in row]
                if not any(row_values):  # 跳过空行
                    continue

                # 构建文本：Header1: Val1, Header2: Val2 ...
                content_parts = []
                for header, val in zip(headers, row_values):
                    if val:
                        content_parts.append(f"{header}: {val}")

                content = " | ".join(content_parts)
                documents.append(
                    Document(
                        page_content=content,
                        metadata={"source": str(file_path), "sheet": sheet_name, "row": row_idx, "type": "excel"},
                    )
                )
        return documents


class WordParser(FileParser):
    async def parse(self, file_path: Path) -> list[Document]:
        """
        Word 解析：提取段落级结构。
        """

        # python-docx 的加载和遍历通常很快，但为了保险依然放入线程
        def _parse_word():
            doc = DocxDocument(file_path)
            results = []
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if text:
                    page_estimate = (i // 20) + 1
                    results.append(
                        Document(
                            page_content=text,
                            metadata={
                                "source": str(file_path),
                                "paragraph_index": i,
                                "page_estimate": page_estimate,
                                "type": "word",
                            },
                        )
                    )
            return results

        return await asyncio.to_thread(_parse_word)


class PPTParser(FileParser):
    async def parse(self, file_path: Path) -> list[Document]:
        """
        PPT 解析：利用 Unstructured 按幻灯片提取。
        """
        loader = UnstructuredPowerPointLoader(str(file_path), mode="elements")
        # Unstructured 加载较慢
        docs = await asyncio.to_thread(loader.load)

        # 转换 metadata 格式
        for doc in docs:
            page = doc.metadata.get("page_number", 1)
            doc.metadata.update({"source": str(file_path), "slide": page, "type": "ppt"})
        return docs
